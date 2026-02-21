"""
CurrHub — AI Curriculum Generator  (Production Backend)
Flask + IBM Granite 3.3 2B via Ollama
"""

from flask import Flask, request, jsonify, Response, stream_with_context
from flask_cors import CORS
import requests, json, time, csv, re
from io import BytesIO, StringIO

# ── Export libraries ──────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.lib import colors

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

MODEL_ID       = "granite3.2:2b"
OLLAMA_URL     = "http://localhost:11434/api/generate"
OLLAMA_CHAT    = "http://localhost:11434/api/chat"
OLLAMA_TIMEOUT = 180   # seconds

print("\n" + "═"*65)
print("  CurrHub — AI Curriculum Generator  (Production)")
print(f"  Model  : {MODEL_ID}")
print(f"  Ollama : {OLLAMA_URL}")
print("═"*65 + "\n")

# ─────────────────────────────────────────────────────────────────────────────
# OLLAMA HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def ollama_generate(prompt: str, timeout: int = OLLAMA_TIMEOUT) -> str | None:
    """Single non-streaming completion via /api/generate"""
    try:
        r = requests.post(
            OLLAMA_URL,
            json={
                "model": MODEL_ID,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.35,
                    "top_p": 0.9,
                    "num_ctx": 4096,
                },
            },
            timeout=timeout,
        )
        if r.status_code == 200:
            return r.json().get("response", "").strip()
        print(f"[Ollama] HTTP {r.status_code}: {r.text[:200]}")
        return None
    except requests.exceptions.ConnectionError:
        print("[Ollama] Connection refused — is `ollama serve` running?")
        return None
    except Exception as exc:
        print(f"[Ollama] Error: {exc}")
        return None


def ollama_chat_stream(messages: list[dict], timeout: int = 120):
    """
    Streaming chat via /api/chat  (fixes the broken chatbot).
    Yields SSE chunks so the React frontend can render tokens as they arrive.
    messages = [{"role": "user"|"assistant"|"system", "content": "..."}]
    """
    try:
        with requests.post(
            OLLAMA_CHAT,
            json={
                "model": MODEL_ID,
                "messages": messages,
                "stream": True,
                "options": {"temperature": 0.5, "top_p": 0.95, "num_ctx": 4096},
            },
            stream=True,
            timeout=timeout,
        ) as resp:
            if resp.status_code != 200:
                yield f"data: {json.dumps({'error': f'Ollama returned {resp.status_code}'})}\n\n"
                return
            for raw_line in resp.iter_lines():
                if not raw_line:
                    continue
                try:
                    chunk = json.loads(raw_line.decode("utf-8"))
                    token = chunk.get("message", {}).get("content", "")
                    done  = chunk.get("done", False)
                    yield f"data: {json.dumps({'token': token, 'done': done})}\n\n"
                    if done:
                        break
                except json.JSONDecodeError:
                    continue
    except requests.exceptions.ConnectionError:
        yield f"data: {json.dumps({'error': 'Cannot connect to Ollama. Run: ollama serve'})}\n\n"
    except Exception as exc:
        yield f"data: {json.dumps({'error': str(exc)})}\n\n"


def parse_ai_json(text: str) -> dict | None:
    """Robust JSON extraction from AI output (handles markdown fences, extra prose)."""
    if not text:
        return None
    # Strip markdown fences
    clean = text.strip()
    for fence in ("```json", "```JSON", "```"):
        if fence in clean:
            parts = clean.split(fence)
            for part in parts:
                part = part.strip().rstrip("`").strip()
                if part.startswith("{"):
                    clean = part
                    break
    # Find outermost JSON object
    start = clean.find("{")
    end   = clean.rfind("}")
    if start != -1 and end != -1 and end > start:
        clean = clean[start : end + 1]
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        # Last attempt: fix common escaping issues
        try:
            return json.loads(re.sub(r"[\x00-\x1f\x7f]", " ", clean))
        except Exception:
            return None


# ─────────────────────────────────────────────────────────────────────────────
# LANGUAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
LANG_SYSTEM = {
    "English": {
        "instruction": "Write ALL text values in English.",
        "example_course": "Introduction to Machine Learning",
        "example_topic":  "Supervised Learning Algorithms",
    },
    "Hindi": {
        "instruction": (
            "CRITICAL: You MUST write ALL course names, semester titles, descriptions, "
            "topics, capstone_project, and learning_outcomes in Hindi using Devanagari script. "
            "Do NOT use any English words in text fields. Use proper Hindi.\n"
            "Example course name : 'मशीन लर्निंग का परिचय'\n"
            "Example topic       : 'पर्यवेक्षित शिक्षण एल्गोरिदम'\n"
            "Example outcome     : 'मशीन लर्निंग की मूल अवधारणाओं में महारत हासिल करें'"
        ),
        "example_course": "मशीन लर्निंग का परिचय",
        "example_topic":  "पर्यवेक्षित शिक्षण",
    },
    "Telugu": {
        "instruction": (
            "CRITICAL: You MUST write ALL course names, semester titles, descriptions, "
            "topics, capstone_project, and learning_outcomes in Telugu script. "
            "Do NOT use any English words in text fields.\n"
            "Example course name : 'మెషిన్ లెర్నింగ్ పరిచయం'\n"
            "Example topic       : 'సూపర్వైజ్డ్ లెర్నింగ్ అల్గారిథమ్స్'\n"
            "Example outcome     : 'మెషిన్ లెర్నింగ్ యొక్క మూల భావాలను నేర్చుకోండి'"
        ),
        "example_course": "మెషిన్ లెర్నింగ్ పరిచయం",
        "example_topic":  "సూపర్వైజ్డ్ లెర్నింగ్",
    },
}

LANG_LABELS = {
    "English": {
        "semester": "Semester", "course": "Course Name", "credits": "Credits",
        "topics": "Topics", "capstone": "Capstone Project",
        "resources": "Learning Resources",
    },
    "Hindi": {
        "semester": "सेमेस्टर", "course": "कोर्स का नाम", "credits": "क्रेडिट",
        "topics": "विषय", "capstone": "कैपस्टोन प्रोजेक्ट",
        "resources": "शिक्षण सामग्री",
    },
    "Telugu": {
        "semester": "సెమిస్టర్", "course": "కోర్సు పేరు", "credits": "క్రెడిట్లు",
        "topics": "విషయాలు", "capstone": "కాప్స్టోన్ ప్రాజెక్ట్",
        "resources": "అభ్యాస వనరులు",
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# RESOURCES DATABASE
# ─────────────────────────────────────────────────────────────────────────────
RESOURCES_DB = {
    "machine learning": {
        "youtube": [
            {"title": "ML by Andrew Ng — Stanford",  "url": "https://www.youtube.com/watch?v=jGwO_UgTS7I"},
            {"title": "ML Full Course — Simplilearn", "url": "https://www.youtube.com/watch?v=NWONeJKn6kc"},
        ],
        "opensource": [
            {"title": "Google ML Crash Course", "url": "https://developers.google.com/machine-learning/crash-course"},
            {"title": "fast.ai",                "url": "https://course.fast.ai/"},
        ],
        "certifications": [
            {"title": "IBM ML Professional Cert",        "url": "https://www.coursera.org/professional-certificates/ibm-machine-learning"},
            {"title": "Stanford ML Specialization",      "url": "https://www.coursera.org/specializations/machine-learning-introduction"},
        ],
    },
    "web development": {
        "youtube": [
            {"title": "Full Stack Web Dev — freeCodeCamp", "url": "https://www.youtube.com/watch?v=zJSY8tbf_ys"},
            {"title": "React Full Course 2024",            "url": "https://www.youtube.com/watch?v=CgkZ7MvWUAA"},
        ],
        "opensource": [
            {"title": "MDN Web Docs",       "url": "https://developer.mozilla.org/"},
            {"title": "The Odin Project",   "url": "https://www.theodinproject.com/"},
        ],
        "certifications": [
            {"title": "Meta Front-End Developer", "url": "https://www.coursera.org/professional-certificates/meta-front-end-developer"},
            {"title": "freeCodeCamp Cert (Free)", "url": "https://www.freecodecamp.org/certification"},
        ],
    },
    "python": {
        "youtube": [
            {"title": "Python Full Course — Mosh",     "url": "https://www.youtube.com/watch?v=_uQrJ0TkZlc"},
            {"title": "Python for Data Science — fCC", "url": "https://www.youtube.com/watch?v=r-uOLxNrNk8"},
        ],
        "opensource": [
            {"title": "Python Official Docs",  "url": "https://docs.python.org/3/"},
            {"title": "Real Python Tutorials", "url": "https://realpython.com/"},
        ],
        "certifications": [
            {"title": "Google IT Automation with Python", "url": "https://grow.google/certificates/it-automation/"},
            {"title": "IBM Python for Data Science",      "url": "https://www.coursera.org/professional-certificates/ibm-data-science"},
        ],
    },
    "data science": {
        "youtube": [
            {"title": "Data Science Full Course — Simplilearn", "url": "https://www.youtube.com/watch?v=ua-CiDNNj30"},
            {"title": "Statistics for DS — StatQuest",          "url": "https://www.youtube.com/watch?v=xxpc-HPKN28"},
        ],
        "opensource": [
            {"title": "Kaggle Learn (Free)",   "url": "https://www.kaggle.com/learn"},
            {"title": "Towards Data Science",  "url": "https://towardsdatascience.com/"},
        ],
        "certifications": [
            {"title": "IBM Data Science Professional Cert", "url": "https://www.coursera.org/professional-certificates/ibm-data-science"},
            {"title": "Google Data Analytics Certificate",  "url": "https://grow.google/certificates/data-analytics/"},
        ],
    },
    "artificial intelligence": {
        "youtube": [
            {"title": "AI For Everyone — Andrew Ng",    "url": "https://www.youtube.com/watch?v=Rb1BDfQnwQA"},
            {"title": "MIT OpenCourseWare AI 6.034",    "url": "https://www.youtube.com/playlist?list=PLUl4u3cNGP63gFHB6xb-kVBiQHYe_4hSi"},
        ],
        "opensource": [
            {"title": "Hugging Face Course", "url": "https://huggingface.co/learn"},
            {"title": "OpenAI Spinning Up",  "url": "https://spinningup.openai.com/"},
        ],
        "certifications": [
            {"title": "IBM AI Engineering Professional",   "url": "https://www.coursera.org/professional-certificates/ai-engineer"},
            {"title": "DeepLearning.AI TensorFlow Dev",   "url": "https://www.coursera.org/professional-certificates/tensorflow-in-practice"},
        ],
    },
}

def get_resources(skill: str) -> dict:
    s = skill.lower()
    for key, val in RESOURCES_DB.items():
        if key in s or any(w in s for w in key.split()):
            return val
    return {
        "youtube":        [{"title": f"Learn {skill} — freeCodeCamp", "url": "https://www.youtube.com/@freecodecamp"}],
        "opensource":     [{"title": "MIT OpenCourseWare",             "url": "https://ocw.mit.edu/"}],
        "certifications": [{"title": "Coursera Free Audit",            "url": "https://www.coursera.org/"}],
    }


# ─────────────────────────────────────────────────────────────────────────────
# CURRICULUM GENERATION — AI PROMPT
# ─────────────────────────────────────────────────────────────────────────────
# ── Prerequisite logic ────────────────────────────────────────────────────────
def assign_prerequisites(semesters: list[dict]) -> list[dict]:
    """
    Post-process: for each semester > 1, pick 1-2 courses from the
    previous semester as prerequisites for some courses in the current one.
    This ensures the Visual Map always has arrows to draw.
    """
    all_codes_by_sem = {}
    for sem in semesters:
        snum = sem.get("semester_number", 1)
        all_codes_by_sem[snum] = [c.get("code", "") for c in sem.get("courses", [])]

    for sem in semesters:
        snum = sem.get("semester_number", 1)
        if snum <= 1:
            continue
        prev_codes = all_codes_by_sem.get(snum - 1, [])
        if not prev_codes:
            continue
        for idx, course in enumerate(sem.get("courses", [])):
            # Only add if not already set
            if course.get("prerequisites"):
                continue
            # Every other course gets a prerequisite from previous semester
            if idx % 2 == 0 and prev_codes:
                course["prerequisites"] = [prev_codes[min(idx, len(prev_codes) - 1)]]
    return semesters


def build_curriculum_prompt(
    skill: str, level: str, semesters: int,
    industry_focus: str, curriculum_type: str, language: str,
) -> str:
    lang_cfg   = LANG_SYSTEM.get(language, LANG_SYSTEM["English"])
    lang_rule  = lang_cfg["instruction"]
    ex_course  = lang_cfg["example_course"]
    ex_topic   = lang_cfg["example_topic"]

    type_note = (
        "This is a SELF-STUDY curriculum — practical, project-based, self-paced. "
        "Each course completable in 4-6 weeks."
        if curriculum_type == "self_study"
        else
        "This is a COLLEGE/ACADEMIC curriculum — include credits, theory, practicals, "
        "prerequisites between courses, and research components."
    )

    # Build sample semester to show exact expected shape
    sample = json.dumps({
        "semester_number": 1,
        "semester_title": ex_course[:15],
        "courses": [
            {
                "code": "CS101",
                "name": ex_course,
                "credits": 4,
                "weekly_hours": 3,
                "description": "One sentence description.",
                "topics": [ex_topic, ex_topic, ex_topic, ex_topic, ex_topic],
                "prerequisites": [],
                "difficulty": "Beginner",
            }
        ],
    }, ensure_ascii=False, indent=2)

    return f"""You are an expert educational curriculum designer.

LANGUAGE RULE — READ CAREFULLY:
{lang_rule}

TASK:
Create a {level}-level curriculum for: {skill}
{type_note}
Industry focus: {industry_focus or 'General Technology'}
Number of semesters: {semesters}
Courses per semester: 3

IMPORTANT CONSTRAINTS:
- ALL text fields (name, description, topics, semester_title, capstone_project, learning_outcomes)
  MUST be written in {language}. Code fields and JSON keys remain in English.
- Example course name in {language}: "{ex_course}"
- Example topic in {language}: "{ex_topic}"
- Each course must have EXACTLY 5 topics
- difficulty must be one of: "Beginner", "Intermediate", "Advanced"
- prerequisites: use course codes from EARLIER semesters only (e.g., ["CS101"])
  Make at least 2 courses have prerequisites from the previous semester.
- Return ONLY valid JSON — no explanation, no markdown fences, no extra text.

REQUIRED JSON FORMAT (example of ONE semester shown):
{sample}

Now return the COMPLETE JSON object with {semesters} semesters:
{{
  "semesters": [ ... {semesters} semester objects ... ],
  "capstone_project": "description in {language}",
  "learning_outcomes": ["outcome 1 in {language}", "outcome 2", "outcome 3"]
}}"""


# ─────────────────────────────────────────────────────────────────────────────
# FALLBACK CURRICULUM (runs if AI returns bad JSON)
# ─────────────────────────────────────────────────────────────────────────────
_KNOWN = {
    "machine learning": [
        ["Supervised Learning", "Unsupervised Learning", "Model Evaluation", "Feature Engineering", "Cross-Validation"],
        ["Neural Networks", "Backpropagation", "Activation Functions", "Regularisation", "TensorFlow Basics"],
        ["Convolutional Neural Nets", "Image Classification", "Transfer Learning", "Object Detection", "ResNet"],
        ["RNN & LSTM", "Sequence Models", "NLP Basics", "Text Classification", "Word Embeddings"],
        ["Reinforcement Learning", "MDP", "Q-Learning", "Policy Gradients", "Actor-Critic"],
        ["MLOps", "Model Deployment", "Docker", "REST API Design", "Monitoring"],
        ["Transformers", "BERT", "GPT Architecture", "Attention Mechanism", "Fine-tuning"],
        ["AI Ethics", "Fairness", "Explainability", "Privacy in ML", "Governance"],
    ],
    "web development": [
        ["HTML5 Fundamentals", "Semantic HTML", "Forms & Validation", "Accessibility", "SEO Basics"],
        ["CSS3 & Flexbox", "CSS Grid", "Animations", "Responsive Design", "CSS Variables"],
        ["JavaScript Basics", "DOM Manipulation", "ES6+ Features", "Async/Await", "Fetch API"],
        ["React Framework", "Components & Props", "State Management", "Hooks", "React Router"],
        ["Node.js & Express", "REST APIs", "Middleware", "JWT Authentication", "Session Handling"],
        ["Databases", "SQL & PostgreSQL", "MongoDB", "ORM with Prisma", "Database Design"],
        ["Deployment & DevOps", "Docker", "CI/CD Pipelines", "Cloud Basics", "Performance"],
        ["Advanced Patterns", "GraphQL", "Microservices", "WebSockets", "Testing with Jest"],
    ],
    "python": [
        ["Python Basics", "Variables & Data Types", "Control Flow", "Functions", "Exception Handling"],
        ["OOP in Python", "Classes & Objects", "Inheritance", "Polymorphism", "Decorators"],
        ["NumPy", "Arrays & Matrices", "Broadcasting", "Linear Algebra Ops", "Performance"],
        ["Pandas", "DataFrames", "Data Cleaning", "GroupBy & Aggregation", "Merging"],
        ["Matplotlib & Seaborn", "Plot Types", "Customisation", "Plotly Interactive", "Dashboards"],
        ["Statistics", "Descriptive Stats", "Probability", "Hypothesis Testing", "Regression"],
        ["Web Scraping", "BeautifulSoup", "Scrapy", "Selenium", "Working with APIs"],
        ["Scikit-learn", "Model Training", "Pipelines", "Feature Selection", "Evaluation"],
    ],
    "data science": [
        ["DS Fundamentals", "Python Basics", "Stats Intro", "EDA", "Data Types"],
        ["Data Collection", "Web APIs", "Scraping", "Databases", "Data Cleaning"],
        ["Exploratory Analysis", "Visualisation", "Correlation", "Outlier Detection", "Feature Analysis"],
        ["ML Basics", "Supervised Learning", "Unsupervised Learning", "Model Selection", "Cross-Validation"],
        ["Advanced ML", "Ensemble Methods", "Gradient Boosting", "XGBoost", "Feature Engineering"],
        ["Deep Learning", "Neural Networks", "CNN", "RNN", "Transfer Learning"],
        ["Big Data", "Apache Spark", "Hadoop Basics", "Cloud Platforms", "Distributed Computing"],
        ["End-to-End Projects", "Pipelines", "Deployment", "Storytelling", "Business Impact"],
    ],
    "artificial intelligence": [
        ["AI Foundations", "History of AI", "Problem Solving", "Search Algorithms", "Knowledge Representation"],
        ["Maths for AI", "Linear Algebra", "Calculus", "Probability", "Optimisation"],
        ["Machine Learning", "Supervised", "Unsupervised", "Reinforcement Learning", "Model Evaluation"],
        ["Deep Learning", "Neural Networks", "Backpropagation", "CNN", "RNN"],
        ["NLP", "Text Processing", "Sentiment Analysis", "Named Entity Recognition", "Transformers"],
        ["Computer Vision", "Image Processing", "Object Detection", "Segmentation", "GANs"],
        ["Advanced RL", "MDPs", "Q-Learning", "Policy Gradients", "Multi-Agent Systems"],
        ["AI Ethics & Deployment", "Bias in AI", "Explainability", "MLOps", "Governance"],
    ],
}

_FALLBACK_PHASES_EN = ["Foundation Phase", "Core Concepts", "Intermediate Level", "Advanced Topics",
                        "Specialisation", "Applied Research", "Industry Projects", "Capstone Phase"]
_FALLBACK_PHASES_HI = ["नींव चरण", "मूल अवधारणाएं", "मध्यवर्ती स्तर", "उन्नत विषय",
                        "विशेषज्ञता", "अनुप्रयुक्त अनुसंधान", "उद्योग परियोजनाएं", "कैपस्टोन चरण"]
_FALLBACK_PHASES_TE = ["పునాది దశ", "మూల భావాలు", "మధ్యస్థ స్థాయి", "అధునాతన అంశాలు",
                        "ప్రత్యేకత", "అనువర్తిత పరిశోధన", "పరిశ్రమ ప్రాజెక్టులు", "కాప్స్టోన్ దశ"]

_PREFIX_EN = ["Introduction to", "Fundamentals of", "Applied", "Advanced",
              "Professional", "Expert", "Deep Dive:", "Mastering"]
_PREFIX_HI = ["परिचय", "मूल सिद्धांत", "अनुप्रयुक्त", "उन्नत",
              "व्यावसायिक", "विशेषज्ञ", "गहन", "मास्टरी"]
_PREFIX_TE = ["పరిచయం", "మూల భావాలు", "అనువర్తిత", "అధునాతన",
              "వృత్తిపరమైన", "నిపుణుడు", "లోతైన", "నైపుణ్యం"]


def build_fallback(
    skill: str, level: str, semesters: int,
    industry_focus: str, language: str,
) -> dict:
    skill_l = skill.lower()
    base_topics = None
    for k, v in _KNOWN.items():
        if k in skill_l or any(w in skill_l for w in k.split()):
            base_topics = v
            break

    if not base_topics:
        base_topics = [
            [f"{skill} Foundations", "Core Terminology", "Tools & Environment Setup", "Basic Principles", "First Project"],
            [f"Core {skill} Concepts", "Fundamental Techniques", "Standard Practices", "Data & Input Handling", "Common Patterns"],
            [f"Intermediate {skill}", "Applied Techniques", "Industry Frameworks", "Performance Tuning", "Error Handling"],
            [f"Advanced {skill}", "Expert Techniques", "Design Patterns", "Scalability", "Security Basics"],
            [f"{skill} & {industry_focus}", "Real-world Applications", "Case Studies", "Integration", "Industry Tools"],
            [f"Professional {skill}", "Team Collaboration", "Code Reviews", "Documentation", "Agile Methods"],
            [f"Specialised {skill}", f"{skill} for {industry_focus}", "Emerging Trends", "Research Methods", "Innovation"],
            [f"{skill} Mastery", "Portfolio Projects", "Open Source Contribution", "Leadership", "Capstone Prep"],
        ]

    phases = (
        _FALLBACK_PHASES_HI if language == "Hindi"
        else _FALLBACK_PHASES_TE if language == "Telugu"
        else _FALLBACK_PHASES_EN
    )
    prefixes = (
        _PREFIX_HI if language == "Hindi"
        else _PREFIX_TE if language == "Telugu"
        else _PREFIX_EN
    )

    diff_map = {0: "Beginner", 1: "Beginner", 2: "Intermediate", 3: "Intermediate",
                4: "Advanced",  5: "Advanced",  6: "Advanced",   7: "Advanced"}

    # Build course code prefix from skill initials
    words  = skill.split()
    prefix = "".join(w[0].upper() for w in words[:2]).ljust(2, "X")[:2]

    sems_out = []
    prev_codes: list[str] = []

    for s_num in range(1, semesters + 1):
        courses = []
        sem_codes: list[str] = []
        phase_idx = min(s_num - 1, len(phases) - 1)

        for c_idx in range(3):
            global_idx  = (s_num - 1) * 3 + c_idx
            topic_idx   = min(global_idx, len(base_topics) - 1)
            px_idx      = min(global_idx, len(prefixes) - 1)
            code        = f"{prefix}{s_num}{c_idx+1:02d}"
            diff        = diff_map.get(s_num - 1, "Intermediate")

            # Course name
            if language == "Hindi":
                name = f"{prefixes[px_idx]} {skill}"
            elif language == "Telugu":
                name = f"{prefixes[px_idx]} {skill}"
            else:
                name = f"{prefixes[px_idx]} {skill.title()}"

            # Description
            if language == "Hindi":
                desc = f"{name} पाठ्यक्रम {industry_focus} उद्योग पर केंद्रित।"
            elif language == "Telugu":
                desc = f"{name} కోర్సు {industry_focus} పరిశ్రమపై దృష్టి పెడుతుంది."
            else:
                desc = f"Comprehensive {skill} course focused on {industry_focus} industry applications."

            # Prerequisites — link to previous semester
            prereqs: list[str] = []
            if s_num > 1 and prev_codes and c_idx % 2 == 0:
                prereqs = [prev_codes[min(c_idx, len(prev_codes) - 1)]]

            courses.append({
                "code": code,
                "name": name,
                "credits": 4,
                "weekly_hours": 3,
                "description": desc,
                "topics": base_topics[topic_idx],
                "prerequisites": prereqs,
                "difficulty": diff,
            })
            sem_codes.append(code)

        sems_out.append({
            "semester_number": s_num,
            "semester_title": phases[phase_idx],
            "courses": courses,
        })
        prev_codes = sem_codes

    # Capstone & outcomes
    if language == "Hindi":
        capstone = f"{skill} का संपूर्ण कैपस्टोन प्रोजेक्ट — {industry_focus} उद्योग के लिए।"
        outcomes = [
            f"{skill} की मूल अवधारणाओं में महारत हासिल करें",
            f"{industry_focus} समस्याओं पर {skill} लागू करें",
            f"पोर्टफोलियो-तैयार परियोजना बनाएं",
        ]
    elif language == "Telugu":
        capstone = f"{skill} సమగ్ర కాప్స్టోన్ ప్రాజెక్ట్ — {industry_focus} పరిశ్రమ కోసం."
        outcomes = [
            f"{skill} యొక్క మూల భావాలను నేర్చుకోండి",
            f"{industry_focus} సమస్యలకు {skill} వర్తించండి",
            f"పోర్ట్‌ఫోలియో ప్రాజెక్ట్ నిర్మించండి",
        ]
    else:
        capstone = f"Complete {skill} Capstone: design, build, deploy, present a real {industry_focus} project."
        outcomes = [
            f"Master core {skill} concepts and tools",
            f"Apply {skill} to real-world {industry_focus} problems",
            f"Build a complete portfolio-ready {skill} project",
        ]

    return {"semesters": sems_out, "capstone_project": capstone, "learning_outcomes": outcomes}


# ─────────────────────────────────────────────────────────────────────────────
# CURRICULUM ASSEMBLY
# ─────────────────────────────────────────────────────────────────────────────
def generate_curriculum(
    skill: str, level: str, semesters: int, weekly_hours: str,
    industry_focus: str, curriculum_type: str, language: str,
) -> dict:

    prompt = build_curriculum_prompt(skill, level, semesters, industry_focus, curriculum_type, language)
    print(f"\n[AI] Generating [{language}] [{curriculum_type}] '{skill}' — {semesters} sems…")
    t0  = time.time()
    raw = ollama_generate(prompt)
    print(f"[AI] Done in {time.time()-t0:.1f}s")

    data = parse_ai_json(raw)
    if data and data.get("semesters") and len(data["semesters"]) > 0:
        print("[AI] JSON parsed successfully — using AI output")
        # Post-process: guarantee prerequisites exist for the visual map
        data["semesters"] = assign_prerequisites(data["semesters"])
    else:
        print("[AI] Bad JSON — using fallback curriculum")
        data = build_fallback(skill, level, semesters, industry_focus, language)

    resources = get_resources(skill) if curriculum_type == "self_study" else {}

    return {
        "skill":                skill,
        "level":                level,
        "weekly_hours":         weekly_hours or "20–25",
        "industry_focus":       industry_focus or "General Technology",
        "curriculum_type":      curriculum_type,
        "language":             language,
        "semesters":            data.get("semesters", []),
        "capstone_project":     data.get("capstone_project", f"Complete {skill} Capstone Project"),
        "learning_outcomes":    data.get("learning_outcomes", []),
        "self_study_resources": resources,
    }


# ─────────────────────────────────────────────────────────────────────────────
# EXPORT ENGINES
# ─────────────────────────────────────────────────────────────────────────────
def export_pdf(curriculum: dict) -> BytesIO:
    buf    = BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=A4,
                               rightMargin=0.6*inch, leftMargin=0.6*inch,
                               topMargin=0.8*inch, bottomMargin=0.8*inch)
    story  = []
    styles = getSampleStyleSheet()
    lang   = curriculum.get("language", "English")
    labels = LANG_LABELS.get(lang, LANG_LABELS["English"])
    PURPLE = colors.HexColor("#7C3AED")
    INDIGO = colors.HexColor("#4F46E5")

    T  = ParagraphStyle("T",  parent=styles["Title"],    fontSize=24, textColor=PURPLE,  spaceAfter=8,  alignment=1, fontName="Helvetica-Bold")
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, textColor=INDIGO,  spaceBefore=16, spaceAfter=6, fontName="Helvetica-Bold")
    M  = ParagraphStyle("M",  parent=styles["Normal"],   fontSize=10, textColor=colors.grey, spaceAfter=10, alignment=1)
    N  = ParagraphStyle("N",  parent=styles["Normal"],   fontSize=9,  spaceAfter=3, leading=14)

    story.append(Paragraph(f"{curriculum.get('skill','').title()} — Learning Plan", T))
    story.append(Paragraph(
        f"Level: {curriculum.get('level','')}  |  Hours/Week: {curriculum.get('weekly_hours','')}  |  "
        f"Industry: {curriculum.get('industry_focus','')}  |  Language: {lang}", M))
    story.append(Spacer(1, 0.15*inch))

    for o in curriculum.get("learning_outcomes", []):
        story.append(Paragraph(f"▸  {o}", N))
    story.append(Spacer(1, 0.1*inch))

    for sem in curriculum.get("semesters", []):
        story.append(Paragraph(f"{labels['semester']} {sem['semester_number']}: {sem.get('semester_title','')}", H2))
        rows = [[labels["course"], "Code", labels["credits"], labels["topics"]]]
        for c in sem.get("courses", []):
            topics = ", ".join(c.get("topics", []))
            if len(topics) > 110: topics = topics[:107] + "…"
            rows.append([c.get("name",""), c.get("code",""), f"{c.get('credits',4)} Cr", topics])
        tbl = Table(rows, colWidths=[2.3*inch, 0.9*inch, 0.7*inch, 3.4*inch])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), PURPLE),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F3FF")]),
            ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#E0E0E0")),
            ("VALIGN",     (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ("LEFTPADDING", (0,0), (-1,-1), 5),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.12*inch))

    story.append(Paragraph(labels["capstone"], H2))
    story.append(Paragraph(curriculum.get("capstone_project", ""), N))

    for cat, items in curriculum.get("self_study_resources", {}).items():
        story.append(Paragraph(cat.replace("_", " ").title(), H2))
        for item in items:
            story.append(Paragraph(f"• {item['title']} — {item['url']}", N))

    doc.build(story)
    buf.seek(0)
    return buf


def export_excel(curriculum: dict) -> BytesIO:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Overview"
    PURPLE = "7C3AED"
    pw = Font(bold=True, color="FFFFFF", size=11)
    pf = PatternFill("solid", fgColor=PURPLE)

    ws["A1"] = f"{curriculum.get('skill','').title()} — Learning Plan"
    ws["A1"].font = Font(bold=True, size=16, color=PURPLE)
    ws.merge_cells("A1:E1")

    for i, (k, v) in enumerate([
        ("Skill",    curriculum.get("skill","")),
        ("Level",    curriculum.get("level","")),
        ("Language", curriculum.get("language","")),
        ("Industry", curriculum.get("industry_focus","")),
        ("Type",     curriculum.get("curriculum_type","").replace("_"," ").title()),
    ], start=3):
        ws.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws.cell(row=i, column=2, value=v)

    for sem in curriculum.get("semesters", []):
        ws_s = wb.create_sheet(title=f"Sem {sem['semester_number']}")
        hdrs = ["Course Name", "Code", "Credits", "Weekly Hrs", "Topics", "Difficulty", "Prerequisites"]
        for col, h in enumerate(hdrs, 1):
            cell = ws_s.cell(row=1, column=col, value=h)
            cell.fill = pf; cell.font = pw
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
        for ri, c in enumerate(sem.get("courses", []), 2):
            alt = PatternFill("solid", fgColor="F5F3FF") if ri%2==0 else PatternFill("solid", fgColor="FFFFFF")
            for ci, val in enumerate([
                c.get("name",""), c.get("code",""), c.get("credits",4), c.get("weekly_hours",3),
                ", ".join(c.get("topics",[])), c.get("difficulty",""),
                ", ".join(c.get("prerequisites",[])),
            ], 1):
                cell = ws_s.cell(row=ri, column=ci, value=val)
                cell.fill = alt; cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws_s.column_dimensions["A"].width = 35
        ws_s.column_dimensions["E"].width = 55

    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return buf


def export_docx(curriculum: dict) -> BytesIO:
    doc  = Document()
    lang = curriculum.get("language","English")
    lbls = LANG_LABELS.get(lang, LANG_LABELS["English"])

    ttl = doc.add_heading(f"{curriculum.get('skill','').title()} — Learning Plan", 0)
    ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ttl.runs: run.font.color.rgb = RGBColor(0x7C,0x3A,0xED)

    meta = doc.add_paragraph()
    meta.add_run(f"Level: {curriculum.get('level','')}   Language: {lang}   Industry: {curriculum.get('industry_focus','')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph()

    for o in curriculum.get("learning_outcomes",[]):
        doc.add_paragraph(o, style="List Bullet")

    for sem in curriculum.get("semesters",[]):
        h = doc.add_heading(f"{lbls['semester']} {sem['semester_number']}: {sem.get('semester_title','')}", level=2)
        for run in h.runs: run.font.color.rgb = RGBColor(0x4F,0x46,0xE5)
        courses = sem.get("courses",[])
        if courses:
            tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
            for ci, ht in enumerate([lbls["course"], "Code", lbls["credits"], lbls["topics"]]):
                cell = tbl.rows[0].cells[ci]; cell.text = ht
                for p in cell.paragraphs:
                    for r in p.runs: r.bold = True
            for c in courses:
                row = tbl.add_row().cells
                row[0].text = c.get("name","")
                row[1].text = c.get("code","")
                row[2].text = f"{c.get('credits',4)} Cr"
                row[3].text = ", ".join(c.get("topics",[]))
        doc.add_paragraph()

    doc.add_heading(lbls["capstone"], level=2)
    doc.add_paragraph(curriculum.get("capstone_project",""))

    for cat, items in curriculum.get("self_study_resources",{}).items():
        doc.add_heading(cat.replace("_"," ").title(), level=3)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item.get("title","")).bold = True
            p.add_run(f" — {item.get('url','')}")

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


def export_markdown(curriculum: dict) -> str:
    lang  = curriculum.get("language","English")
    lbls  = LANG_LABELS.get(lang, LANG_LABELS["English"])
    lines = [f"# {curriculum.get('skill','').title()} — Learning Plan\n"]
    lines.append(f"**Level:** {curriculum.get('level','')} | **Language:** {lang} | **Industry:** {curriculum.get('industry_focus','')}\n")
    for o in curriculum.get("learning_outcomes",[]):
        lines.append(f"- {o}")
    for sem in curriculum.get("semesters",[]):
        lines.append(f"\n## {lbls['semester']} {sem['semester_number']}: {sem.get('semester_title','')}\n")
        lines.append(f"| {lbls['course']} | Code | {lbls['credits']} | {lbls['topics']} |")
        lines.append("|---|---|---|---|")
        for c in sem.get("courses",[]):
            lines.append(f"| {c.get('name','')} | {c.get('code','')} | {c.get('credits',4)} Cr | {', '.join(c.get('topics',[]))} |")
    lines.append(f"\n## {lbls['capstone']}\n{curriculum.get('capstone_project','')}")
    return "\n".join(lines)


def export_csv(curriculum: dict) -> str:
    output = StringIO()
    w = csv.writer(output)
    w.writerow(["Semester","Title","Course Name","Code","Credits","Topics","Difficulty","Prerequisites"])
    for sem in curriculum.get("semesters",[]):
        for c in sem.get("courses",[]):
            w.writerow([
                sem.get("semester_number",""), sem.get("semester_title",""),
                c.get("name",""), c.get("code",""), c.get("credits",4),
                " | ".join(c.get("topics",[])), c.get("difficulty",""),
                ", ".join(c.get("prerequisites",[])),
            ])
    output.seek(0)
    return output.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────────────────────────────────────
@app.route("/api/generate-curriculum", methods=["POST","OPTIONS"])
def route_generate():
    if request.method == "OPTIONS":
        return jsonify({"status":"ok"}), 200
    try:
        d              = request.json
        skill          = d.get("skill","").strip()
        level          = d.get("level","").strip()
        sems           = d.get("semesters")
        weekly_hours   = d.get("weekly_hours","").strip()
        industry_focus = d.get("industry_focus","").strip()
        curriculum_type= d.get("curriculum_type","college")
        language       = d.get("language","English")

        if not skill or not level:
            return jsonify({"error": "skill and level are required"}), 400
        try:
            sems = int(sems)
            if not 2 <= sems <= 8: raise ValueError
        except (TypeError, ValueError):
            return jsonify({"error": "semesters must be 2–8"}), 400

        c = generate_curriculum(skill, level, sems, weekly_hours, industry_focus, curriculum_type, language)
        return jsonify({"success": True, "curriculum": c})
    except Exception as exc:
        print(f"[Generate] Error: {exc}")
        return jsonify({"error": str(exc)}), 500


@app.route("/api/chat-refine", methods=["POST","OPTIONS"])
def route_chat_refine():
    """
    Fixed chatbot endpoint — uses Ollama /api/chat for proper multi-turn chat.
    Supports both streaming (SSE) and non-streaming JSON.
    The React frontend should call with stream=true in the body for SSE.
    """
    if request.method == "OPTIONS":
        return jsonify({"status":"ok"}), 200

    try:
        d          = request.json
        curriculum = d.get("curriculum", {})
        message    = d.get("message","").strip()
        history    = d.get("history", [])   # [{role, content}, ...]
        use_stream = d.get("stream", False)

        if not message:
            return jsonify({"error": "message is required"}), 400

        # Build system message with curriculum context
        sems  = curriculum.get("semesters", [])
        c_list = [c["name"] for sem in sems for c in sem.get("courses", [])]
        system_msg = (
            f"You are an expert educational curriculum assistant for CurrHub.\n"
            f"The user has generated a curriculum for: {curriculum.get('skill','unknown skill')}\n"
            f"Level: {curriculum.get('level','')}, Semesters: {len(sems)}, "
            f"Language: {curriculum.get('language','English')}\n"
            f"Courses: {', '.join(c_list[:12])}{'...' if len(c_list)>12 else ''}\n\n"
            f"Help the user refine this curriculum. Give specific, actionable suggestions. "
            f"Keep responses under 250 words. End with: 'Shall I help you with anything else?'"
        )

        # Build message list for Ollama chat
        ollama_msgs = [{"role": "system", "content": system_msg}]
        # Add prior conversation history
        for turn in history:
            ollama_msgs.append({"role": turn.get("role","user"), "content": turn.get("content","")})
        # Add current user message
        ollama_msgs.append({"role": "user", "content": message})

        if use_stream:
            def generate_sse():
                yield f"data: {json.dumps({'token': '', 'done': False})}\n\n"
                for chunk in ollama_chat_stream(ollama_msgs):
                    yield chunk
            return Response(
                stream_with_context(generate_sse()),
                mimetype="text/event-stream",
                headers={
                    "Cache-Control": "no-cache",
                    "X-Accel-Buffering": "no",
                },
            )
        else:
            # Non-streaming fallback (collect all tokens)
            full_response = ""
            for chunk_str in ollama_chat_stream(ollama_msgs):
                # Parse the SSE data line
                if chunk_str.startswith("data: "):
                    try:
                        chunk_data = json.loads(chunk_str[6:])
                        if "error" in chunk_data:
                            return jsonify({"error": chunk_data["error"]}), 500
                        full_response += chunk_data.get("token", "")
                    except json.JSONDecodeError:
                        pass

            if not full_response.strip():
                full_response = (
                    "I couldn't generate a response right now. "
                    "Please check that Ollama is running with: ollama serve"
                )
            return jsonify({"success": True, "response": full_response.strip()})

    except Exception as exc:
        print(f"[Chat] Error: {exc}")
        return jsonify({"error": str(exc)}), 500


@app.route("/api/export", methods=["POST","OPTIONS"])
def route_export():
    if request.method == "OPTIONS":
        return jsonify({"status":"ok"}), 200
    try:
        import base64
        d          = request.json
        curriculum = d.get("curriculum", {})
        fmt        = d.get("format","pdf").lower()
        name       = curriculum.get("skill","curriculum").replace(" ","_")

        if not curriculum:
            return jsonify({"error":"No curriculum data provided"}), 400

        if fmt == "pdf":
            buf = export_pdf(curriculum)
            enc = base64.b64encode(buf.getvalue()).decode()
            return jsonify({"success":True,"data":enc,"filename":f"{name}_curriculum.pdf","mime":"application/pdf"})
        elif fmt == "excel":
            buf = export_excel(curriculum)
            enc = base64.b64encode(buf.getvalue()).decode()
            return jsonify({"success":True,"data":enc,"filename":f"{name}_curriculum.xlsx",
                            "mime":"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"})
        elif fmt == "docx":
            buf = export_docx(curriculum)
            enc = base64.b64encode(buf.getvalue()).decode()
            return jsonify({"success":True,"data":enc,"filename":f"{name}_curriculum.docx",
                            "mime":"application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
        elif fmt == "markdown":
            md  = export_markdown(curriculum)
            enc = base64.b64encode(md.encode("utf-8")).decode()
            return jsonify({"success":True,"data":enc,"filename":f"{name}_curriculum.md","mime":"text/markdown"})
        elif fmt == "csv":
            csv_text = export_csv(curriculum)
            enc = base64.b64encode(csv_text.encode("utf-8")).decode()
            return jsonify({"success":True,"data":enc,"filename":f"{name}_curriculum.csv","mime":"text/csv"})
        elif fmt == "json":
            j   = json.dumps(curriculum, indent=2, ensure_ascii=False)
            enc = base64.b64encode(j.encode("utf-8")).decode()
            return jsonify({"success":True,"data":enc,"filename":f"{name}_curriculum.json","mime":"application/json"})
        else:
            return jsonify({"error":f"Unsupported format: {fmt}"}), 400

    except Exception as exc:
        print(f"[Export] Error: {exc}")
        return jsonify({"error": str(exc)}), 500


@app.route("/health", methods=["GET"])
def health():
    # Check Ollama connectivity
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=4)
        ollama_ok = r.status_code == 200
        models    = [m["name"] for m in r.json().get("models",[])]
    except Exception:
        ollama_ok = False
        models    = []
    return jsonify({
        "status":        "healthy",
        "model":         MODEL_ID,
        "ollama_online": ollama_ok,
        "available_models": models,
        "timestamp":     time.time(),
    })


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"  Visit : http://localhost:5000")
    print(f"  Health: http://localhost:5000/health")
    print(f"  Ensure Ollama is running: ollama serve\n")
    app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
